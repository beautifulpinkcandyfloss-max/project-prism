"""
Generic readers: take a file path or URL, return raw text (or bytes for
binary formats). These know NOTHING about Ra, Seth, sessions, etc. -- that
logic belongs in parsers/. Keeping this dumb-and-generic means the same
reader code works for all six sources.
"""

import io
import requests
import pdfplumber
import docx
from bs4 import BeautifulSoup


def is_url(value: str) -> bool:
    return value.lower().startswith("http://") or value.lower().startswith("https://")


def read_txt_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx_file(path: str) -> str:
    doc = docx.Document(path)
    return '\n'.join(p.text for p in doc.paragraphs)


def read_docx_bytes(content: bytes) -> str:
    doc = docx.Document(io.BytesIO(content))
    return '\n'.join(p.text for p in doc.paragraphs)


def read_pdf_file(path: str, x_tolerance: float = 0.5) -> str:
    """x_tolerance controls how large a horizontal gap between characters
    has to be before pdfplumber treats it as a word boundary. The
    library's default (3) assumes fairly generously-spaced text; some
    PDFs encode genuine inter-word gaps much tighter than that, and the
    default then merges entire sentences into one run-on string with no
    spaces at all (confirmed by reproducing it against a synthetic PDF
    with tight word spacing -- the default setting merged everything,
    while a lower value recovered correct spacing). If you still see
    merged words after this, try an even lower value (e.g. 0.5)."""
    pages = []
    skipped = 0
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                t = page.extract_text(x_tolerance=x_tolerance)
            except Exception as exc:
                print(f"[PDF] Skipping page {i + 1} of '{path}': "
                      f"{type(exc).__name__}: {exc}")
                skipped += 1
                continue
            if t:
                pages.append(t)
    if skipped:
        print(f"[PDF] '{path}': extracted {len(pages)} pages, "
              f"skipped {skipped} unreadable page(s)")
    return '\n'.join(pages)


def read_pdf_bytes(content: bytes, x_tolerance: float = 0.5) -> str:
    pages = []
    skipped = 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                t = page.extract_text(x_tolerance=x_tolerance)
            except Exception as exc:
                print(f"[PDF] Skipping page {i + 1}: {type(exc).__name__}: {exc}")
                skipped += 1
                continue
            if t:
                pages.append(t)
    if skipped:
        print(f"[PDF] extracted {len(pages)} pages, skipped {skipped} unreadable page(s)")
    return '\n'.join(pages)


def fetch_url_bytes(url: str) -> bytes:
    response = requests.get(url, timeout=20)
    response.raise_for_status()
    return response.content


def read_url_html(url: str) -> str:
    """Fetch a URL and return visible text with line breaks preserved
    between block-level elements -- important so parsers can still find
    'SESSION 12' or a date sitting on its own line."""
    response = requests.get(url, timeout=20)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    lines = [s for s in soup.stripped_strings]
    return '\n'.join(lines)


def read_source_list(list_path: str):
    """Read a docs.txt / pdfs.txt / urls.txt manifest file -> list of
    non-empty lines (local filenames or URLs)."""
    items = []
    try:
        with open(list_path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                candidate = line.strip()
                if candidate and not candidate.startswith('#'):
                    items.append(candidate)
    except FileNotFoundError:
        pass
    return items


def read_any(source: str, base_dir: str) -> str:
    """Read a doc/pdf source whether it's a local filename or a URL.
    Returns raw text (decoded for docx/pdf, raw HTML-stripped text for URLs)."""
    import os
    lower = source.lower()

    if is_url(source):
        if lower.endswith(".docx"):
            return read_docx_bytes(fetch_url_bytes(source))
        if lower.endswith(".pdf"):
            return read_pdf_bytes(fetch_url_bytes(source))
        return read_url_html(source)

    path = source if os.path.isabs(source) else os.path.join(base_dir, source)
    if lower.endswith(".docx"):
        return read_docx_file(path)
    if lower.endswith(".pdf"):
        return read_pdf_file(path)
    return read_txt_file(path)
